from docx import Document
from io import BytesIO

# Load your DOCX file (uploaded_file_bytes comes from your upload widget)
doc = Document(BytesIO(uploaded_file_bytes))

superscripts = []

# Find all superscript numbers in document
for para in doc.paragraphs:
    for run in para.runs:
        # Check if run is superscript AND is a digit
        if run.font.superscript and run.text.strip().isdigit():
            superscripts.append((run.text.strip(), para.text.strip()))

# Find paragraphs that start with a matching reference number
def find_reference(number):
    for para in doc.paragraphs:
        if para.text.strip().startswith(number + '. '):
            return para.text.strip()
    return None

# Pair superscript with reference paragraph
references = []
for num, sup_text in superscripts:
    ref_text = find_reference(num)
    if ref_text:
        references.append((num, sup_text, ref_text))

print("Superscripts found:", superscripts)
print("References paired:", references)
